#!/usr/bin/env python3
"""
TensorVault - Backend Service
Private business knowledge assistant. Indexes your documents into a local
vector store, answers questions grounded in your docs, exports answers as
PDF or DOCX. Runs entirely on your own machine.

Architecture (v0.1.0):
  - LLM        : Gemma 4 via Ollama  (GPU or CPU, whichever Ollama detects)
  - Embeddings : mxbai-embed-large via Ollama  (same GPU as LLM)
  - Index      : FAISS in-memory (Flat IP, cosine after L2 normalization)
  - Storage    : SQLite for document chunks
  - STT        : faster-whisper (small, auto-detects CUDA via ctranslate2)
  - Files      : pdfplumber, python-docx for text extraction
  - Export     : fpdf2 for PDF, python-docx for DOCX

Target hardware: 4+ core CPU, 8 GB RAM (16 GB recommended). GPU is optional
but strongly recommended for usable LLM speed.

This service bundles to ~200 MB after PyInstaller compression (no PyTorch,
no sentence-transformers, no transformers). The full installer including
Electron + Ollama binary sits well under the 2 GB GitHub release limit.
"""

import os
import sys
import json

# Fix Unicode output on Windows (PyInstaller console uses CP1252 by default)
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try: sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    except Exception: pass
if sys.stderr and hasattr(sys.stderr, 'reconfigure'):
    try: sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception: pass

import io
import sqlite3
import time
import tempfile
import shutil
import subprocess
import threading
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Generator

import numpy as np
import faiss
import requests as http_requests
from flask import Flask, request, jsonify, Response, stream_with_context, send_file
from flask_cors import CORS

# Note: TensorVault uses Ollama for BOTH the LLM and embeddings. We do not
# bundle PyTorch / sentence-transformers in the .exe (saves ~1 GB of wheels
# and keeps the installer under the 2 GB GitHub asset cap). Embeddings run
# on the same GPU Ollama uses for the LLM.

# -- Hardware target ----------------------------------------------------------
CPU_THREADS = int(os.environ.get('TENSORVAULT_CPU_THREADS', '4'))
os.environ['OMP_NUM_THREADS']        = str(CPU_THREADS)
os.environ['MKL_NUM_THREADS']        = str(CPU_THREADS)
os.environ['OPENBLAS_NUM_THREADS']   = str(CPU_THREADS)
os.environ['NUMEXPR_NUM_THREADS']    = str(CPU_THREADS)
faiss.omp_set_num_threads(CPU_THREADS)

# -- Paths --------------------------------------------------------------------
if getattr(sys, 'frozen', False):
    BASE_DIR = Path(sys.executable).parent
else:
    BASE_DIR = Path(__file__).parent

USER_DATA_DIR = Path(os.environ.get('TENSORVAULT_USER_DIR',
                     Path.home() / 'AppData' / 'Roaming' / 'TensorVault' / 'user_docs'))

# User document index
USER_FAISS  = USER_DATA_DIR / 'user.faiss'
USER_DB     = USER_DATA_DIR / 'user_chunks.db'
USER_META   = USER_DATA_DIR / 'user_meta.json'

# Export directory (where generated PDFs/DOCX go before download)
EXPORT_DIR  = USER_DATA_DIR.parent / 'exports'
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# -- Service config -----------------------------------------------------------
HOST            = os.environ.get('RAG_HOST', '127.0.0.1')
PORT            = int(os.environ.get('RAG_PORT', '8712'))

# Ollama (LLM + embeddings — both run on whatever device Ollama detects).
OLLAMA_HOST        = os.environ.get('OLLAMA_HOST',         'http://127.0.0.1:11434')
OLLAMA_MODEL       = os.environ.get('OLLAMA_MODEL',        'gemma4')
OLLAMA_EMBED_MODEL = os.environ.get('OLLAMA_EMBED_MODEL',  'mxbai-embed-large')

# Whisper STT — bundled via faster-whisper. Auto-detects CUDA if installed.
WHISPER_MODEL   = os.environ.get('WHISPER_MODEL',  'small')
def _detect_whisper_device():
    # We dropped torch from the bundle, so detect CUDA via ctranslate2 instead.
    try:
        from ctranslate2 import get_supported_compute_types
        return 'cuda' if 'float16' in get_supported_compute_types('cuda') else 'cpu'
    except Exception:
        return 'cpu'
WHISPER_DEVICE  = _detect_whisper_device()
WHISPER_COMPUTE = 'float16' if WHISPER_DEVICE == 'cuda' else 'int8'

# Piper TTS - CPU, zero VRAM
if getattr(sys, 'frozen', False):
    _PIPER_DEFAULT = str(BASE_DIR / 'piper' / 'piper.exe')
    _PIPER_VOICE_DEFAULT = str(BASE_DIR / 'piper' / 'en_US-amy-medium.onnx')
else:
    _PIPER_DEFAULT = shutil.which('piper') or 'piper'
    _PIPER_VOICE_DEFAULT = str(BASE_DIR / 'piper_voices' / 'en_US-amy-medium.onnx')
PIPER_BIN       = os.environ.get('PIPER_BIN', _PIPER_DEFAULT)
PIPER_VOICE     = os.environ.get('PIPER_VOICE', _PIPER_VOICE_DEFAULT)

# OCR binaries - tesseract + poppler bundled alongside service.exe
if getattr(sys, 'frozen', False):
    _tesseract_dir = BASE_DIR / 'tesseract'
    _poppler_dir = BASE_DIR / 'poppler' / 'Library' / 'bin'
    if _tesseract_dir.exists():
        os.environ.setdefault('TESSDATA_PREFIX', str(_tesseract_dir / 'tessdata'))
        os.environ['PATH'] = str(_tesseract_dir) + os.pathsep + os.environ['PATH']
    if _poppler_dir.exists():
        os.environ['PATH'] = str(_poppler_dir) + os.pathsep + os.environ['PATH']

# RAG config
RAG_CTX_CHUNKS  = int(os.environ.get('RAG_CTX_CHUNKS', '5'))
RAG_CTX_CHARS   = int(os.environ.get('RAG_CTX_CHARS',  '900'))
CHUNK_SIZE      = 400
CHUNK_OVERLAP   = 50

# -- Globals ------------------------------------------------------------------
# Embedding dimension is discovered from the first embedding call (different
# Ollama embedding models have different dimensions: 384, 768, 1024, ...).
embed_dim: Optional[int] = None

user_index  = None
user_db_con = None
user_meta:  List[Dict] = []
user_id_map: List[int] = []   # maps FAISS position -> SQLite row id

whisper_model = None
_whisper_lock = threading.Lock()
ollama_process = None

app = Flask(__name__)
CORS(app)

# Per-thread SQLite connections (SQLite connections aren't thread-safe)
_user_local = threading.local()


def get_user_con() -> Optional[sqlite3.Connection]:
    if not USER_DB.exists():
        return None
    if not hasattr(_user_local, 'con'):
        con = sqlite3.connect(str(USER_DB), check_same_thread=False)
        con.execute("PRAGMA cache_size=-16384")
        con.execute("PRAGMA temp_store=MEMORY")
        con.row_factory = sqlite3.Row
        _user_local.con = con
    return _user_local.con


# -- Startup ------------------------------------------------------------------
def init():
    print('=' * 64)
    print('TensorVault - Private Business Knowledge Assistant')
    print('Requires: 4+ core CPU, 8 GB RAM (16 GB recommended)')
    print('GPU optional but recommended (NVIDIA 4+ GB VRAM)')
    print('=' * 64)
    print(f'  User dir       : {USER_DATA_DIR}')
    print(f'  CPU threads    : {CPU_THREADS}')
    print(f'  LLM model      : {OLLAMA_MODEL}  (via Ollama)')
    print(f'  Embedder model : {OLLAMA_EMBED_MODEL}  (via Ollama)')
    print(f'  Whisper device : {WHISPER_DEVICE}')
    print()

    USER_DATA_DIR.mkdir(parents=True, exist_ok=True)

    # Load user index if exists. We do not load any heavy ML models here —
    # embeddings happen via Ollama on demand. This keeps backend startup
    # near-instant.
    _load_user_index()

    print('TensorVault ready -> http://{}:{}'.format(HOST, PORT))
    print('-' * 64)


# -- User index ---------------------------------------------------------------
def _open_user_db(write=False):
    USER_DATA_DIR.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(str(USER_DB))
    con.execute("PRAGMA journal_mode=WAL")
    con.execute("PRAGMA synchronous=NORMAL")
    con.execute("""
        CREATE TABLE IF NOT EXISTS chunks (
            id    INTEGER PRIMARY KEY AUTOINCREMENT,
            name  TEXT NOT NULL,
            title TEXT NOT NULL,
            text  TEXT NOT NULL
        )
    """)
    con.commit()
    con.row_factory = sqlite3.Row
    return con


def _load_user_index():
    global user_index, user_meta, user_id_map
    if USER_FAISS.exists():
        try:
            user_index = faiss.read_index(str(USER_FAISS))
            print(f'  User index: {user_index.ntotal} vectors')
        except Exception as e:
            print(f'  User index load failed: {e}')
    if USER_META.exists():
        user_meta = json.loads(USER_META.read_text())
    # Rebuild FAISS -> SQLite id map from DB
    if USER_DB.exists() and user_index is not None:
        con = _open_user_db()
        user_id_map = [r['id'] for r in con.execute("SELECT id FROM chunks ORDER BY id").fetchall()]
        con.close()


def _save_user_index():
    if user_index and user_index.ntotal > 0:
        faiss.write_index(user_index, str(USER_FAISS))
    USER_META.write_text(json.dumps(user_meta, indent=2))


# -- Embedding (via Ollama) ---------------------------------------------------
class EmbedError(RuntimeError):
    """Raised when Ollama's embedding endpoint fails after retries."""


def _embed_one(text: str, attempts: int = 3) -> List[float]:
    """Embed a single string. Retries with backoff for the first-call lag
    that happens when Ollama is loading the embedding model into VRAM."""
    last_err = None
    delay = 1.0
    for i in range(attempts):
        try:
            r = http_requests.post(
                f'{OLLAMA_HOST}/api/embeddings',
                json={'model': OLLAMA_EMBED_MODEL, 'prompt': text},
                timeout=120,
            )
            r.raise_for_status()
            v = r.json().get('embedding') or []
            if v:
                return v
            last_err = 'empty embedding returned'
        except Exception as e:
            last_err = str(e)
        if i < attempts - 1:
            time.sleep(delay)
            delay *= 2  # 1s, 2s, 4s
    raise EmbedError(f'Ollama embedding failed after {attempts} tries: {last_err}')


def embed(texts: List[str]) -> np.ndarray:
    """Embed texts via Ollama's /api/embeddings endpoint.

    Ollama runs the embedding model on whichever device it detected (GPU if
    available, CPU otherwise). No PyTorch dependency.

    Raises EmbedError if Ollama is unreachable or returns no embedding so
    the caller can surface a real failure to the user instead of silently
    indexing zero-dim garbage.
    """
    global embed_dim
    if not texts:
        return np.zeros((0, embed_dim or 1), dtype='float32')
    vecs = []
    for text in texts:
        v = _embed_one(text)   # raises on persistent failure
        vecs.append(v)
    arr = np.array(vecs, dtype='float32')
    # L2-normalize so we can use FAISS inner-product as cosine similarity.
    norms = np.linalg.norm(arr, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    arr = arr / norms
    if embed_dim is None:
        embed_dim = arr.shape[1]
    return arr


# -- Search -------------------------------------------------------------------
def search_user(qvec: np.ndarray, k: int) -> List[Dict]:
    global user_index
    if user_index is None or user_index.ntotal == 0:
        return []
    k = min(k, user_index.ntotal)
    distances, indices = user_index.search(qvec, k)
    con = get_user_con()
    if con is None:
        return []
    faiss_to_sql = {}
    for i in indices[0]:
        if 0 <= i < len(user_id_map):
            faiss_to_sql[int(i)] = user_id_map[int(i)]
    sql_ids = list(faiss_to_sql.values())
    if not sql_ids:
        return []
    rows = {r['id']: (r['title'], r['text'], r['name'])
            for r in con.execute(
                f"SELECT id, title, text, name FROM chunks WHERE id IN ({','.join('?'*len(sql_ids))})",
                sql_ids
            ).fetchall()}
    hits = []
    for dist, idx in zip(distances[0], indices[0]):
        idx = int(idx)
        sql_id = faiss_to_sql.get(idx)
        if sql_id is None or sql_id not in rows:
            continue
        title, text, name = rows[sql_id]
        hits.append({
            'title': title, 'text': text, 'doc': name,
            'score': float(dist), 'source': 'user', 'chunk_id': sql_id,
        })
    return hits


def rerank(query: str, hits: List[Dict], top_k: int) -> List[Dict]:
    """No-op reranker for v0.1.0.

    The previous version used a cross-encoder (sentence-transformers) for
    second-pass relevance scoring. That dragged in ~1 GB of PyTorch wheels.
    For v0.1.0 we trust the top-k FAISS scores from a high-quality embedder
    (mxbai-embed-large or equivalent). A v0.2 enhancement can add LLM-based
    reranking via Ollama if quality wins out.
    """
    return hits[:top_k]


# -- Text extraction for user docs --------------------------------------------
def extract_text(path: str) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    if ext in ('.txt', '.md', '.csv', '.rtf'):
        return p.read_text(encoding='utf-8', errors='replace')
    if ext == '.pdf':
        # Try native text extraction first (fast)
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(p)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t and t.strip():
                        pages.append(t)
            if pages:
                return '\n\n'.join(pages)
        except ImportError:
            pass
        # Fallback: OCR for scanned/image PDFs - process page-by-page to avoid OOM
        try:
            from pdf2image import convert_from_path
            import pytesseract
            import pymupdf
            total_pages = len(pymupdf.open(str(p)))
            print(f'  [OCR] Scanned PDF detected - {total_pages} pages, running OCR on {p.name}...',
                  flush=True)
            import gc
            pages = []
            for i in range(1, total_pages + 1):
                imgs = convert_from_path(str(p), dpi=200, first_page=i, last_page=i,
                                         thread_count=4)
                if imgs:
                    t = pytesseract.image_to_string(imgs[0])
                    if t and t.strip():
                        pages.append(t)
                    imgs[0].close()
                    del imgs
                    gc.collect()
                if i % 20 == 0:
                    print(f'  [OCR] {i}/{total_pages} pages...', flush=True)
            print(f'  [OCR] Done - {len(pages)} pages with text', flush=True)
            if pages:
                return '\n\n'.join(pages)
        except ImportError:
            raise RuntimeError('pip install pdfplumber pdf2image pytesseract')
        raise ValueError('No text found - PDF may be empty or unsupported')
    if ext in ('.docx', '.doc'):
        try:
            import docx
            return '\n'.join(para.text for para in docx.Document(str(p)).paragraphs)
        except ImportError:
            raise RuntimeError('pip install python-docx')
    return p.read_text(encoding='utf-8', errors='replace')


def chunk_text(text: str, title: str) -> List[Dict]:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunks.append({'title': title, 'text': ' '.join(words[i:i + CHUNK_SIZE])})
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


# -- Ollama helpers -----------------------------------------------------------
def start_ollama_if_needed():
    global ollama_process
    try:
        r = http_requests.get(f'{OLLAMA_HOST}/api/tags', timeout=3)
        if r.status_code == 200:
            print('  Ollama: already running')
            return
    except Exception:
        pass

    candidates = [
        str(BASE_DIR / 'ollama' / 'ollama.exe'),
        str(BASE_DIR / 'ollama' / 'ollama'),
        shutil.which('ollama') or '',
    ]
    ollama_bin = next((c for c in candidates if c and Path(c).exists()), 'ollama')

    env = {
        **os.environ,
        'OLLAMA_MODELS': str(BASE_DIR / 'ollama_models'),
        'CUDA_VISIBLE_DEVICES': '0',
    }
    print(f'  Ollama: starting ({ollama_bin})...')
    try:
        ollama_process = subprocess.Popen(
            [ollama_bin, 'serve'],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            env=env,
        )
        for _ in range(30):
            time.sleep(0.5)
            try:
                if http_requests.get(f'{OLLAMA_HOST}/api/tags', timeout=2).status_code == 200:
                    print('  Ollama: ready')
                    return
            except Exception:
                pass
        print('  Ollama: WARNING - slow to start')
    except FileNotFoundError:
        print(f'  Ollama: not found at {ollama_bin}')


def _pull_one_model(name: str):
    """Pull a single Ollama model, no-op if already present."""
    try:
        r = http_requests.get(f'{OLLAMA_HOST}/api/tags', timeout=5)
        names = [m.get('name', '') for m in r.json().get('models', [])]
        base  = name.split(':')[0]
        if any(base in n for n in names):
            print(f'  Model {name}: already present')
            return
    except Exception:
        pass

    print(f'  Pulling {name} (first run only)...')
    try:
        with http_requests.post(
            f'{OLLAMA_HOST}/api/pull',
            json={'name': name},
            stream=True, timeout=3600,
        ) as r:
            for line in r.iter_lines():
                if not line:
                    continue
                d = json.loads(line)
                status = d.get('status', '')
                total  = d.get('total', 0)
                done   = d.get('completed', 0)
                pct    = done / total * 100 if total else 0
                if 'pulling' in status or 'verifying' in status:
                    print(f'  {status} {pct:.1f}%', end='\r', flush=True)
        print(f'\n  Model {name}: ready')
    except Exception as e:
        print(f'  WARNING: pull failed for {name}: {e}')


def ensure_model_pulled():
    """Pull both the LLM and the embedding model."""
    _pull_one_model(OLLAMA_MODEL)
    _pull_one_model(OLLAMA_EMBED_MODEL)


# -- Whisper STT --------------------------------------------------------------
def get_whisper():
    global whisper_model
    with _whisper_lock:
        if whisper_model is None:
            try:
                from faster_whisper import WhisperModel
                print(f'  Loading Whisper {WHISPER_MODEL} on {WHISPER_DEVICE}...')
                whisper_model = WhisperModel(
                    WHISPER_MODEL,
                    device=WHISPER_DEVICE,
                    compute_type=WHISPER_COMPUTE,
                )
                print('  Whisper: ready')
            except ImportError:
                print('  Whisper: faster-whisper not installed')
    return whisper_model


# -- LLM prompt builder -------------------------------------------------------
SYSTEM_PROMPT = (
    "You are TensorVault, a private business knowledge assistant. The user "
    "has indexed their own business documents (contracts, policies, "
    "playbooks, meeting notes, RFPs, SOPs, internal wikis, sales collateral) "
    "and is asking you questions about them.\n"
    "\n"
    "Core rules:\n"
    "1. Ground every claim in the indexed documents. Cite each fact with "
    "[N] matching the numbered context passages above. Lead with the answer.\n"
    "2. If multiple documents address the same point, cite each and note "
    "any discrepancies plainly.\n"
    "3. If the documents do not contain the answer, say so directly. Do not "
    "guess. You may offer general knowledge but mark it clearly as outside "
    "the indexed sources.\n"
    "4. Never invent: dates, dollar amounts, party names, signing dates, "
    "deadlines, jurisdictions, policy numbers, percentages, or any other "
    "specific identifier. If it is not in the documents, say so.\n"
    "\n"
    "Format the response to the request:\n"
    "- 'Summarize ...' -> produce a focused summary, not a Q&A.\n"
    "- 'Compare ...' -> use a clear side-by-side or bulleted structure.\n"
    "- 'Draft ...' -> produce the requested artifact (email, memo, clause).\n"
    "- 'List' or 'extract' -> produce a clean list with citations.\n"
    "- Otherwise -> direct prose, answer first.\n"
    "\n"
    "Tone: professional, concise, no filler. Match the formality of the "
    "indexed documents. No greetings or sign-offs unless drafting one."
)


def build_rag_prompt(question: str, hits: List[Dict]) -> str:
    if hits:
        ctx = '\n\n'.join(
            f"[{i}] (Doc: {h.get('doc', h.get('title',''))}) {h.get('title','')}\n"
            f"{h.get('text','')[:RAG_CTX_CHARS]}"
            for i, h in enumerate(hits, 1)
        )
        return (
            f"{SYSTEM_PROMPT}\n\n"
            f"CONTEXT (numbered passages from the user's indexed documents):\n"
            f"{ctx}\n\n"
            f"QUESTION: {question}\n\n"
            f"Answer the question using the context passages above. Cite "
            f"each fact with [1], [2], etc."
        )
    else:
        return (
            f"{SYSTEM_PROMPT}\n\n"
            f"CONTEXT: No relevant passages were found in the user's "
            f"indexed documents.\n\n"
            f"QUESTION: {question}\n\n"
            f"State directly that no document sources matched. You may "
            f"offer general knowledge if it helps, but make clear it is "
            f"not drawn from the user's documents."
        )


def stream_ollama(prompt: str) -> Generator[str, None, None]:
    payload = {
        'model':  OLLAMA_MODEL,
        'prompt': prompt,
        'stream': True,
        'think':  True,
        'options': {
            'temperature':  0.6,
            'top_p':        0.95,
            'num_ctx':      4096,
            'num_thread':   CPU_THREADS,
        },
    }
    try:
        with http_requests.post(
            f'{OLLAMA_HOST}/api/generate',
            json=payload, stream=True, timeout=300,
        ) as resp:
            if resp.status_code != 200:
                yield f'data: {json.dumps({"type":"error","message":f"Ollama {resp.status_code}"})}\n\n'
                return
            in_think = False
            for line in resp.iter_lines():
                if not line:
                    continue
                try:
                    chunk = json.loads(line)
                except Exception:
                    continue
                think_tok = chunk.get('thinking', '')
                ans_tok   = chunk.get('response', '')
                if think_tok:
                    yield f'data: {json.dumps({"type":"thinking","token":think_tok})}\n\n'
                if ans_tok:
                    if '<think>' in ans_tok:
                        in_think = True
                    if in_think:
                        yield f'data: {json.dumps({"type":"thinking","token":ans_tok})}\n\n'
                        if '</think>' in ans_tok:
                            in_think = False
                    else:
                        yield f'data: {json.dumps({"type":"answer","token":ans_tok})}\n\n'
                if chunk.get('done'):
                    yield f'data: {json.dumps({"type":"done"})}\n\n'
                    return
    except Exception as e:
        yield f'data: {json.dumps({"type":"error","message":str(e)})}\n\n'


# -- Export helpers (PDF / DOCX) ----------------------------------------------
def _strip_tags(text: str) -> str:
    """Remove [[u]]/[[g]]/[[n]] source tags from the LLM output for clean export."""
    import re
    return re.sub(r'\[\[/?[ugn]\]\]', '', text)


def make_docx(title: str, body: str, sources: Optional[List[Dict]] = None) -> bytes:
    """Generate a DOCX file from a title + body + optional sources list."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    h = doc.add_heading(title or 'TensorVault Report', level=1)
    for para in _strip_tags(body).split('\n\n'):
        p = doc.add_paragraph(para.strip())
        for run in p.runs:
            run.font.size = Pt(11)
    if sources:
        doc.add_heading('Sources', level=2)
        for i, s in enumerate(sources, 1):
            doc.add_paragraph(
                f"[{i}] {s.get('doc', s.get('title','source'))} - "
                f"{s.get('title','')}",
                style='List Number',
            )
    doc.add_paragraph()
    footer = doc.add_paragraph()
    fr = footer.add_run('Generated by TensorVault')
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def make_pdf(title: str, body: str, sources: Optional[List[Dict]] = None) -> bytes:
    """Generate a PDF file from a title + body + optional sources list."""
    from fpdf import FPDF
    pdf = FPDF(format='Letter', unit='mm')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    # Title
    pdf.set_font('Helvetica', 'B', 16)
    pdf.multi_cell(0, 10, title or 'TensorVault Report')
    pdf.ln(2)
    # Body
    pdf.set_font('Helvetica', '', 11)
    clean = _strip_tags(body)
    # FPDF needs latin-1 safe text; replace unicode that doesn't fit
    safe = clean.encode('latin-1', errors='replace').decode('latin-1')
    for para in safe.split('\n\n'):
        pdf.multi_cell(0, 6, para.strip())
        pdf.ln(2)
    # Sources
    if sources:
        pdf.ln(4)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Sources', ln=True)
        pdf.set_font('Helvetica', '', 10)
        for i, s in enumerate(sources, 1):
            line = f"[{i}] {s.get('doc', s.get('title','source'))} - {s.get('title','')}"
            safe_line = line.encode('latin-1', errors='replace').decode('latin-1')
            pdf.multi_cell(0, 5, safe_line)
    # Footer
    pdf.ln(6)
    pdf.set_font('Helvetica', 'I', 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, 'Generated by TensorVault', ln=True)
    out = pdf.output(dest='S')
    if isinstance(out, str):
        out = out.encode('latin-1')
    return bytes(out)


# -- Routes -------------------------------------------------------------------
@app.route('/health')
def health():
    return jsonify({
        'status':           'healthy',
        'app':              'TensorVault',
        'version':          '0.1.0',
        'user_docs':        len(user_meta),
        'user_chunks':      user_index.ntotal if user_index else 0,
        'whisper_device':   WHISPER_DEVICE,
        'llm_model':        OLLAMA_MODEL,
        'embed_model':      OLLAMA_EMBED_MODEL,
    })


@app.route('/ollama/status')
def ollama_status():
    try:
        r      = http_requests.get(f'{OLLAMA_HOST}/api/tags', timeout=3)
        models = [m['name'] for m in r.json().get('models', [])]
        ready  = any(OLLAMA_MODEL.split(':')[0] in m for m in models)
        return jsonify({'running': True, 'models': models, 'model_ready': ready})
    except Exception:
        return jsonify({'running': False, 'models': [], 'model_ready': False})


@app.route('/search', methods=['POST', 'GET'])
def search():
    if request.method == 'GET':
        q = request.args.get('q', '').strip()
        k = min(int(request.args.get('k', 5)), 20)
    else:
        d = request.json or {}
        q = d.get('q', '').strip()
        k = min(int(d.get('k', 5)), 20)

    if not q:
        return jsonify({'error': 'No query'}), 400

    t0 = time.time()
    try:
        qvec        = embed([q])
        candidate_k = max(k * 5, 50)
        hits = search_user(qvec, candidate_k)
        hits = rerank(q, hits, k)
        for h in hits:
            h['text'] = h['text'][:600]
        return jsonify({
            'query': q, 'results': hits,
            'num_results': len(hits),
            'search_time_ms': round((time.time() - t0) * 1000, 1),
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/ask', methods=['POST'])
def ask():
    d        = request.json or {}
    question = d.get('q', '').strip()
    k        = min(int(d.get('k', RAG_CTX_CHUNKS)), 10)

    if not question:
        return jsonify({'error': 'No question'}), 400

    try:
        qvec        = embed([question])
        candidate_k = max(k * 4, 20)
        hits = search_user(qvec, candidate_k)
        hits = rerank(question, hits, k)
    except Exception as e:
        return jsonify({'error': f'RAG: {e}'}), 500

    prompt  = build_rag_prompt(question, hits)
    sources_payload = [
        {'title': h.get('title',''), 'doc': h.get('doc',''),
         'text': h.get('text','')[:200]}
        for h in hits
    ]

    def event_stream():
        yield f'data: {json.dumps({"type":"sources","sources":sources_payload})}\n\n'
        yield from stream_ollama(prompt)

    return Response(
        stream_with_context(event_stream()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/transcribe', methods=['POST'])
def transcribe():
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio'}), 400
    w = get_whisper()
    if w is None:
        return jsonify({'error': 'Whisper unavailable'}), 503
    f   = request.files['audio']
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.webm')
    try:
        f.save(tmp.name); tmp.close()
        segs, info = w.transcribe(tmp.name, beam_size=5)
        text = ' '.join(s.text for s in segs).strip()
        return jsonify({'text': text, 'language': info.language,
                        'duration_s': round(info.duration, 2)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try: os.unlink(tmp.name)
        except Exception: pass


@app.route('/speak', methods=['POST'])
def speak():
    text = (request.json or {}).get('text', '').strip()
    if not text:
        return jsonify({'error': 'No text'}), 400
    if not (Path(PIPER_BIN).exists() or shutil.which('piper')):
        return jsonify({'error': 'Piper not installed'}), 503
    if not Path(PIPER_VOICE).exists():
        return jsonify({'error': f'Voice model not found: {PIPER_VOICE}'}), 503
    try:
        import wave
        proc = subprocess.run(
            [PIPER_BIN, '--model', PIPER_VOICE, '--output_raw'],
            input=text.encode(), capture_output=True, timeout=30,
        )
        if proc.returncode != 0:
            raise RuntimeError(proc.stderr.decode())
        buf = io.BytesIO()
        with wave.open(buf, 'wb') as wf:
            wf.setnchannels(1); wf.setsampwidth(2); wf.setframerate(22050)
            wf.writeframes(proc.stdout)
        buf.seek(0)
        return Response(buf.read(), mimetype='audio/wav')
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# -- Document ingestion -------------------------------------------------------
@app.route('/ingest', methods=['POST'])
def ingest_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f   = request.files['file']
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=Path(f.filename).suffix)
    try:
        f.save(tmp.name); tmp.close()
        return jsonify(_ingest_file(tmp.name, f.filename))
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try: os.unlink(tmp.name)
        except Exception: pass


@app.route('/ingest_path', methods=['POST'])
def ingest_path():
    path = (request.json or {}).get('path', '').strip()
    if not path or not os.path.isfile(path):
        return jsonify({'error': 'File not found'}), 400
    try:
        return jsonify(_ingest_file(path, Path(path).name))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _ingest_file(path: str, display_name: str) -> Dict:
    global user_index, user_id_map
    text = extract_text(path)
    if not text.strip():
        raise ValueError('No text found')
    title  = Path(display_name).stem.replace('_', ' ').replace('-', ' ')
    chunks = chunk_text(text, title)
    if not chunks:
        raise ValueError('No chunks produced')

    texts = [c['text'] for c in chunks]
    vecs  = embed(texts)
    dim   = vecs.shape[1]

    if user_index is None or user_index.d != dim:
        user_index = faiss.IndexFlatIP(dim)
    user_index.add(vecs)

    con = _open_user_db(write=True)
    con.executemany(
        "INSERT INTO chunks(name, title, text) VALUES (?,?,?)",
        [(display_name, c['title'], c['text']) for c in chunks]
    )
    con.commit()
    all_ids = [r['id'] for r in con.execute("SELECT id FROM chunks ORDER BY id").fetchall()]
    con.close()
    user_id_map = all_ids

    if hasattr(_user_local, 'con'):
        del _user_local.con

    user_meta[:] = [m for m in user_meta if m['name'] != display_name]
    user_meta.append({
        'name':   display_name,
        'chunks': len(chunks),
        'added':  time.strftime('%Y-%m-%d %H:%M:%S'),
    })
    _save_user_index()
    return {'ok': True, 'name': display_name, 'chunks': len(chunks)}


@app.route('/docs')
def docs_list():
    return jsonify({'docs': user_meta})


@app.route('/docs', methods=['DELETE'])
def docs_clear():
    global user_index, user_id_map
    user_index = None
    user_id_map = []
    user_meta.clear()
    for f in [USER_FAISS, USER_DB, USER_META]:
        if Path(str(f)).exists():
            Path(str(f)).unlink()
    if hasattr(_user_local, 'con'):
        del _user_local.con
    return jsonify({'ok': True})


@app.route('/docs/<name>', methods=['DELETE'])
def docs_delete(name):
    global user_index, user_id_map
    con = _open_user_db(write=True)
    con.execute("DELETE FROM chunks WHERE name = ?", (name,))
    con.commit()
    remaining = con.execute("SELECT id, title, text FROM chunks ORDER BY id").fetchall()
    con.close()
    if hasattr(_user_local, 'con'):
        del _user_local.con

    user_meta[:] = [m for m in user_meta if m['name'] != name]

    if remaining:
        texts = [r['text'] for r in remaining]
        vecs  = embed(texts)
        user_index = faiss.IndexFlatIP(vecs.shape[1])
        user_index.add(vecs)
        user_id_map = [r['id'] for r in remaining]
    else:
        user_index = None
        user_id_map = []

    _save_user_index()
    return jsonify({'ok': True})


# -- Export endpoints (NEW in TensorVault) ------------------------------------
@app.route('/export/docx', methods=['POST'])
def export_docx():
    """Generate a DOCX from a title + body + optional sources list."""
    d = request.json or {}
    title   = d.get('title', 'TensorVault Report')
    body    = d.get('body', '')
    sources = d.get('sources', None)
    if not body.strip():
        return jsonify({'error': 'No body content'}), 400
    try:
        data = make_docx(title, body, sources)
        safe_name = ''.join(c for c in title if c.isalnum() or c in ' _-').strip() or 'report'
        return send_file(
            io.BytesIO(data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{safe_name}.docx',
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    """Generate a PDF from a title + body + optional sources list."""
    d = request.json or {}
    title   = d.get('title', 'TensorVault Report')
    body    = d.get('body', '')
    sources = d.get('sources', None)
    if not body.strip():
        return jsonify({'error': 'No body content'}), 400
    try:
        data = make_pdf(title, body, sources)
        safe_name = ''.join(c for c in title if c.isalnum() or c in ' _-').strip() or 'report'
        return send_file(
            io.BytesIO(data),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{safe_name}.pdf',
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# -- Entry point --------------------------------------------------------------
if __name__ == '__main__':
    init()
    threading.Thread(
        target=lambda: (start_ollama_if_needed(), ensure_model_pulled()),
        daemon=True,
    ).start()
    app.run(host=HOST, port=PORT, debug=False, threaded=True)
